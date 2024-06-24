import streamlit as st
from langchain_google_genai import GoogleGenerativeAI
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from PyPDF2 import PdfReader
import docx
from pydantic import BaseModel, Field

template = """
Kindly help me to set {number} questions and coresponding answers based on this provided context

context:{context}
format_instructions: {format_instructions}
"""

def format_context(docs):
    return "\n".join(doc.page_content for doc in docs)

class QA_parser(BaseModel):
    question: str = Field("Questions generated")
    answer: str = Field("Answer for each of the question")
    
    
class LambdaStreamlitLoader(BaseLoader):
    def __init__(self, file) -> None:
        self.file = file

    def lazy_load(self):
        file_name = file.name
        *_, ext = file_name.split(".")

        if ext == "docx" or ext == "doc":
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                for line in paragraph.text.split("\n"):
                    yield Document(page_content=line)

        elif ext == "pdf":
            doc = PdfReader(file)
            for page in doc.pages:
                for line in page.extract_text().split("\n"):
                    yield Document(page_content=line)
                    
help_template = """
I have been asked this question: {question}
And then I provided this answer: {student_answer}
However, this is the actual answer: {ground_truth}
You are not to respond with the actual answer, just compare my answer to the right answer and guide me
If my answer indicates that I don't know, guide me with a clearer step by step approach to make me understand it, with words of encouragement
The answer doesn't have to be verbatim, tell me I am correct if I use other words that have the same meaning
"""

help_llm = GoogleGenerativeAI(model="gemini-pro")

if "qa_json" not in st.session_state:
    st.session_state.qa_json = None
    
if "count" not in st.session_state:
    st.session_state.count = 0
    

with st.sidebar:
    file = st.file_uploader("Upload document here", type=["pdf", "docx", "doc"])
    
    if file:
        loader = LambdaStreamlitLoader(file)
        document = loader.load()
        st.success("Document Loaded sucessfully")
        number_ip = st.number_input("How many questions would you like to set?", step=1, min_value=1, format="%i")
        
        if st.button("Generate"):
            with st.spinner("Loading..."):
                llm = GoogleGenerativeAI(model="gemini-pro")
                parser = JsonOutputParser()
                full_prompt = template.format(context=format_context(document), number=number_ip, format_instructions=parser.get_format_instructions())

                llm_response = llm(full_prompt)
                qas = parser.parse(llm_response)
                
                st.session_state.qa_json = qas
                st.success("Question and answer ready")
            
          
st.title("Question and Answer session")
if st.session_state.qa_json:
    qa_session = st.session_state.qa_json["questions"]
    try:
        current_qa = qa_session[st.session_state.count]
        current_q = current_qa["question"]
        current_a = current_qa["answer"]
        st.markdown(f"#### Question {st.session_state.count + 1}")
        st.write(f"Question: {current_q}")
        
        user_answer = st.text_area("Provide your answer here", key=st.session_state.count)
        if st.button("Evaluate with AI"):
            with st.spinner("Evaluating"):
                help_template = help_template.format(question=current_q, student_answer=user_answer, ground_truth=current_a)
                help_response = help_llm(help_template)
                st.write(help_response)
        
        if st.button("Reveal answer"):
            st.write("Actual answer from material")
            st.write(current_a)
            
        if st.button("NEXT"):
            if user_answer:
                st.session_state.count += 1
                st.rerun()
            else:
                st.error("You have not provided any answer, you can just say you don't know, you know")
            
    except IndexError:
        st.markdown("## Congratulations, You have come to the end of this session, you are free to regenerate more questions and practice more.")
        st.session_state.count = 0
        
    except Exception as e:
        st.error(f"This error {e} just occured, regenerate new questions please, as this model is a free and unstable one")
    
else:
    st.warning("Upload document and specify number of questions to proceed")


