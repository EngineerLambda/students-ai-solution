import streamlit as st
from langchain.chains.llm import LLMChain
from PyPDF2 import PdfReader
from fpdf import FPDF
import docx
from typing import Iterator
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents.stuff import StuffDocumentsChain
from langchain.prompts import PromptTemplate
from langchain_google_genai import GoogleGenerativeAI

# app config
st.set_page_config(
    page_title="Summarization App", layout="centered", page_icon="📙"
)
st.title("Summarization App")

template = """
You are an AI summarization agent, using the provided document below
Document: {document}
Make sure to preserve headings and titles.
Generate summary for EVERY SINGLE THING discussed in the document
Look to summarize per section or subtopic discussed. Sections in the sense of 1.o, 1.1, 1.3, 2.0, 2.1 and so on
Also make sure the summaries generated are NOT TOO CONCISE, they should include every aspect, I don't want to miss out on any information
If there are types, classes and kinds of things stated, be sure to inlcude them all, if explanations, definitions or descriptions are available for types\
classes and kinds mentioned, make sure to highlight the explanation, definition or description along with them more preferable if you use a semi colon, like below:
The types of feedbacks are positive feedback and negative feedback. postive feedback is so and so and negative feedback is so and so;
Summmarize as
**Types of feedback**
- positive feedback: so and so about positive feedback
- negative feedback: so and so about negative feedback


Add the end of the summary generation, also generate practive questions with heading "PRACTICE QUESTIONS" for user from each of the headings in the summary.
"""


class PDF(FPDF):
    def header(self):
        self.set_font("Arial", "B", 12)

    def footer(self):
        self.set_y(-15)
        self.set_font("Arial", "I", 8)
        self.cell(0, 10, f"Page {self.page_no()}", 0, 0, "C")

    def chapter_body(self, body):
        self.set_font("Arial", "", 12)
        self.multi_cell(0, 10, body)
        self.ln()


def generate_pdf(text, output_name):
    pdf = PDF()
    pdf.add_page()
    pdf.chapter_body(text)
    pdf.output(output_name)


class LambdaStreamlitLoader(BaseLoader):
    def __init__(self, file) -> None:
        self.file = file

    def lazy_load(self) -> Iterator[Document]:
        file_name = file.name
        *_, ext = file_name.split(".")

        if ext == "docx" or ext == "doc":
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                for line in paragraph.text.split("\n"):
                    yield Document(page_content=line)

        elif ext == "pdf":
            doc = PdfReader(file)
            for page in doc.pages:
                for line in page.extract_text().split("\n"):
                    yield Document(page_content=line)


file = st.file_uploader("Upload document here", type=["pdf", "docx", "doc"])

if file:
    loader = LambdaStreamlitLoader(file)
    document = loader.load()
    document_split = RecursiveCharacterTextSplitter(
        separators=["\n", "\n\n"]
    ).split_documents(document)

    *name, ext = file.name.split(".")
    name = "".join(name)
    out_name = "{} [SUMMARY-LAMBDA].{}".format(name, ext)

    if st.button("SUMMARIZE"):
        with st.spinner("Generating Summary ..."):
            prompt = PromptTemplate.from_template(template)
            llm = GoogleGenerativeAI(model="gemini-pro")

            llm_chain = LLMChain(llm=llm, prompt=prompt)
            stuff_chain = StuffDocumentsChain(
                llm_chain=llm_chain, document_variable_name="document"
            )

            response = stuff_chain.run(document_split)
            st.write(response)

        if ext == "doc":
            mime = "application/msword"
        elif ext == "docx":
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext == "pdf":
            generate_pdf(response, output_name=out_name)
            with open(out_name, "rb") as pdf_file:
                response = pdf_file.read()
            mime = "application/octec-stream"

        download = st.download_button(
            label="Download Summary",
            data=response,
            file_name=out_name,
            key="download_summary",
            mime=mime,
        )
