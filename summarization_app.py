from typing import Iterator
import streamlit as st
from langchain.chains.llm import LLMChain
from PyPDF2 import PdfReader
import docx
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains.combine_documents.stuff import StuffDocumentsChain
from langchain.prompts import PromptTemplate
from langchain_google_genai import GoogleGenerativeAI

template = """
You are an AI summarization agent, you are provided with these texts
{text}
Make sure to preserve headings and titles.
And you are to provide bullet points kind of summary, dont miss any important information, explain important concepts without missing any of them.
When there are definitions, use an entire bullet point for that, retain technical or domain specific words too.
And make sure nothing worth noting is missed out but make sure to still not provide irrelevant messages that would not benefit the reader.
If there are types, classes and kinds of things stated, be sure to inlcude them all, if explanations, definitions or descriptions are available for types, classes and kinds mentioned, make sure to highlight the explanation, definition or description along with them more preferable if you use a semi colon, like below:
The types of feedbacks are positive feedback and negative feedback. postive feedback is so and so and negative feedback is so and so;
Summmarize as
**Types of feedback**
- positive feedback: so and so about positive feedback
- negative feedback: so and so about negative feedback
Remember to only do this if it is included in the given material
"""


class LambdaStreamlitLoader(BaseLoader):
    def __init__(self, file) -> None:
        self.file = file

    def lazy_load(self) -> Iterator[Document]:
        file_name = file.name
        *_, ext = file_name.split(".")

        if ext == "docx" or ext == "doc":
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                for line in paragraph.text.split("\n"):
                    yield Document(page_content=line)

        elif ext == "pdf":
            doc = PdfReader(file)
            for page in doc.pages:
                for line in page.extract_text().split("\n"):
                    yield Document(page_content=line)


file = st.file_uploader("Upload document here", type=["pdf", "docx", "doc"])

if file:
    loader = LambdaStreamlitLoader(file)
    document = loader.load()
    document_split = RecursiveCharacterTextSplitter(
        separators=["\n", "\n\n"]
    ).split_documents(document)

    *name, ext = file.name.split(".")
    name = "".join(name)
    out_name = "{} [SUMMARY-LAMBDA].{}".format(name, ext)

    if st.button("SUMMARIZE"):
        with st.spinner("Generating Summary ..."):
            prompt = PromptTemplate.from_template(template)
            llm = GoogleGenerativeAI(model="gemini-pro")

            llm_chain = LLMChain(llm=llm, prompt=prompt)
            stuff_chain = StuffDocumentsChain(
                llm_chain=llm_chain, document_variable_name="text"
            )

            response = stuff_chain.run(document_split)
            st.write(response)

        if ext == "doc":
            mime = "application/msword"
        elif ext == "docx":
            mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext == "pdf":
            mime = "application/pdf"

        download = st.download_button(
            label="Download Summary",
            data=response,
            file_name=out_name,
            key="download_summary",
            mime=mime,
        )
